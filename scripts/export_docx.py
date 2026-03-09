#!/usr/bin/env python3
import argparse
from pathlib import Path
from docx import Document


def main():
    p = argparse.ArgumentParser(description='将 Markdown/TXT 摘要导出为 docx')
    p.add_argument('input', help='输入 markdown/txt 文件')
    p.add_argument('--output', required=True, help='输出 docx 路径')
    args = p.parse_args()

    src = Path(args.input).expanduser().resolve()
    out = Path(args.output).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for line in src.read_text(encoding='utf-8').splitlines():
        if not line.strip():
            doc.add_paragraph('')
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    doc.save(out)
    print(out)


if __name__ == '__main__':
    main()
