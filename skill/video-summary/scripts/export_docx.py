#!/usr/bin/env python3
import argparse
import re
from pathlib import Path


def parse_markdown(content: str) -> list:
    """解析 Markdown 内容，返回段落列表"""
    lines = content.splitlines()
    paragraphs = []
    
    for line in lines:
        line = line.rstrip()
        if not line.strip():
            paragraphs.append({'type': 'blank'})
            continue
        
        # 标题
        if line.startswith('### '):
            paragraphs.append({'type': 'h3', 'text': line[4:]})
        elif line.startswith('## '):
            paragraphs.append({'type': 'h2', 'text': line[3:]})
        elif line.startswith('# '):
            paragraphs.append({'type': 'h1', 'text': line[2:]})
        # 表格（简化处理：检测 | 开头和结尾的行）
        elif line.startswith('|') and '|' in line[1:]:
            paragraphs.append({'type': 'table_row', 'text': line})
        # 列表项
        elif line.startswith('- [x] ') or line.startswith('- [X] '):
            paragraphs.append({'type': 'checkbox_checked', 'text': line[6:]})
        elif line.startswith('- [ ] '):
            paragraphs.append({'type': 'checkbox_unchecked', 'text': line[5:]})
        elif line.startswith('- ') or line.startswith('* '):
            paragraphs.append({'type': 'bullet', 'text': line[2:]})
        elif re.match(r'^\d+\. ', line):
            match = re.match(r'^(\d+)\. (.*)', line)
            paragraphs.append({'type': 'number', 'num': match.group(1), 'text': match.group(2)})
        # 分割线
        elif line.startswith('---') or line.startswith('***') or line.startswith('___'):
            paragraphs.append({'type': 'hr'})
        # 普通段落
        else:
            paragraphs.append({'type': 'p', 'text': line})
    
    return paragraphs


def export_docx(input_path: Path, output_path: Path):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ModuleNotFoundError as exc:
        raise SystemExit("缺少依赖 python-docx，请先安装：pip install python-docx") from exc

    content = input_path.read_text(encoding='utf-8')
    paragraphs = parse_markdown(content)
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        
        if p['type'] == 'blank':
            doc.add_paragraph('')
        
        elif p['type'] == 'h1':
            heading = doc.add_heading(p['text'], level=1)
            heading.style.font.size = Pt(18)
            heading.style.font.bold = True
        
        elif p['type'] == 'h2':
            heading = doc.add_heading(p['text'], level=2)
            heading.style.font.size = Pt(15)
            heading.style.font.bold = True
        
        elif p['type'] == 'h3':
            heading = doc.add_heading(p['text'], level=3)
            heading.style.font.size = Pt(13)
            heading.style.font.bold = True
        
        elif p['type'] == 'hr':
            doc.add_paragraph().add_run().add_break()
        
        elif p['type'] == 'table_row':
            # 收集连续的表格行
            table_rows = [p]
            j = i + 1
            while j < len(paragraphs) and paragraphs[j]['type'] == 'table_row':
                table_rows.append(paragraphs[j])
                j += 1
            
            # 跳过表头分隔符行
            table_data = [row['text'].split('|')[1:-1] for row in table_rows if not re.match(r'^[\s|:/-]+$', row['text'].split('|')[1] if len(row['text'].split('|')) > 1 else '')]
            table_data = [row['text'].split('|')[1:-1] for row in table_rows if not all(re.match(r'^[\s|:/-]+$', cell.strip()) for cell in row['text'].split('|')[1:-1])]
            
            if table_data:
                cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for r_idx, row_data in enumerate(table_data):
                    cells = row_data
                    for c_idx, cell_text in enumerate(cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text.strip()
            
            i = j - 1  # 重置循环索引
        
        elif p['type'] == 'checkbox_checked':
            para = doc.add_paragraph()
            para.add_run('☑ ' + p['text'])
        
        elif p['type'] == 'checkbox_unchecked':
            para = doc.add_paragraph()
            para.add_run('☐ ' + p['text'])
        
        elif p['type'] == 'bullet':
            doc.add_paragraph(p['text'], style='List Bullet')
        
        elif p['type'] == 'number':
            para = doc.add_paragraph()
            para.add_run(f"{p['num']}. {p['text']}")
        
        elif p['type'] == 'p':
            doc.add_paragraph(p['text'])
        
        i += 1
    
    doc.save(output_path)
    print(output_path)


def main():
    p = argparse.ArgumentParser(description='将 Markdown/TXT 摘要导出为 docx（支持更好排版）')
    p.add_argument('input', help='输入 markdown/txt 文件')
    p.add_argument('--output', required=True, help='输出 docx 路径')
    args = p.parse_args()

    src = Path(args.input).expanduser().resolve()
    out = Path(args.output).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    if not src.exists():
        raise SystemExit(f"文件不存在: {src}")

    export_docx(src, out)


if __name__ == '__main__':
    main()
